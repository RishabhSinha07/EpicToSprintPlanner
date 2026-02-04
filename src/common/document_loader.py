"""
Document loader for various file formats.
Supports PDF, DOCX, Markdown, and plain text.
"""
import io
import base64
from typing import BinaryIO, Optional, Tuple, List, Dict
from pathlib import Path


def load_document(file_obj: BinaryIO, file_extension: str, extract_images: bool = True) -> Tuple[str, List[Dict]]:
    """
    Load and extract text and images from various document formats.

    Args:
        file_obj: File-like object containing the document
        file_extension: File extension (e.g., 'pdf', 'docx', 'md', 'txt')
        extract_images: Whether to extract images from the document

    Returns:
        Tuple of (text content, list of image dictionaries)
    """
    file_extension = file_extension.lower().lstrip('.')

    if file_extension == 'pdf':
        return _load_pdf(file_obj, extract_images)
    elif file_extension in ['docx', 'doc']:
        return _load_docx(file_obj, extract_images)
    elif file_extension in ['md', 'markdown']:
        return _load_text(file_obj), []
    elif file_extension == 'txt':
        return _load_text(file_obj), []
    else:
        raise ValueError(f"Unsupported file format: {file_extension}")


def _load_pdf(file_obj: BinaryIO, extract_images: bool = True) -> Tuple[str, List[Dict]]:
    """Extract text and images from PDF file."""
    try:
        import PyPDF2
    except ImportError:
        raise ImportError("PyPDF2 is required for PDF processing. Install with: pip install PyPDF2")

    reader = PyPDF2.PdfReader(file_obj)
    text_parts = []

    for page_num, page in enumerate(reader.pages, 1):
        text = page.extract_text()
        if text.strip():
            text_parts.append(f"--- Page {page_num} ---\n{text}")

    text = "\n\n".join(text_parts)

    images = []
    if extract_images:
        try:
            images = _extract_images_from_pdf(file_obj)
        except Exception as e:
            print(f"Warning: Failed to extract images from PDF: {str(e)}")

    return text, images


def _load_docx(file_obj: BinaryIO, extract_images: bool = True) -> Tuple[str, List[Dict]]:
    """Extract text and images from DOCX file."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required for DOCX processing. Install with: pip install python-docx")

    doc = Document(file_obj)
    text_parts = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = ' | '.join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                text_parts.append(row_text)

    text = "\n\n".join(text_parts)

    images = []
    if extract_images:
        try:
            images = _extract_images_from_docx(doc)
        except Exception as e:
            print(f"Warning: Failed to extract images from DOCX: {str(e)}")

    return text, images


def _load_text(file_obj: BinaryIO) -> str:
    """Load plain text or markdown file."""
    content = file_obj.read()

    # Try UTF-8 first, fall back to latin-1
    try:
        return content.decode('utf-8')
    except UnicodeDecodeError:
        return content.decode('latin-1')


def get_file_extension(filename: str) -> str:
    """Extract file extension from filename."""
    return Path(filename).suffix.lstrip('.')


def _extract_images_from_pdf(file_obj: BinaryIO) -> List[Dict]:
    """
    Extract images from PDF using PyMuPDF.

    Args:
        file_obj: PDF file object

    Returns:
        List of image dictionaries with metadata
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise ImportError("PyMuPDF is required for PDF image extraction. Install with: pip install PyMuPDF")

    # Reset file pointer
    file_obj.seek(0)
    pdf_bytes = file_obj.read()
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")

    images = []
    image_counter = 0

    for page_num in range(len(doc)):
        page = doc[page_num]
        image_list = page.get_images()

        for img_index, img in enumerate(image_list):
            xref = img[0]
            try:
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                image_ext = base_image["ext"]

                # Process image for Bedrock
                processed_bytes, media_type = process_image_for_bedrock(image_bytes, image_ext)

                images.append({
                    "image_id": f"img_{image_counter}",
                    "image_data": processed_bytes,
                    "media_type": media_type,
                    "page_number": page_num + 1,
                    "original_ext": image_ext
                })
                image_counter += 1

            except Exception as e:
                print(f"Warning: Failed to extract image {img_index} from page {page_num + 1}: {str(e)}")
                continue

    doc.close()
    return images


def _extract_images_from_docx(doc) -> List[Dict]:
    """
    Extract images from DOCX document with estimated position tracking.

    Args:
        doc: python-docx Document object

    Returns:
        List of image dictionaries with metadata including estimated position
    """
    images = []
    image_counter = 0

    # Get total document length for position estimation
    total_length = sum(len(para.text) + 2 for para in doc.paragraphs)

    # Extract images from document parts
    for rel_id, rel in doc.part.rels.items():
        if "image" in rel.target_ref:
            try:
                image_part = rel.target_part
                image_bytes = image_part.blob

                # Determine image format from content type
                content_type = image_part.content_type
                if 'jpeg' in content_type or 'jpg' in content_type:
                    ext = 'jpeg'
                elif 'png' in content_type:
                    ext = 'png'
                elif 'gif' in content_type:
                    ext = 'gif'
                elif 'webp' in content_type:
                    ext = 'webp'
                else:
                    ext = 'png'  # Default

                # Process image for Bedrock
                processed_bytes, media_type = process_image_for_bedrock(image_bytes, ext)

                images.append({
                    "image_id": f"img_{image_counter}",
                    "image_data": processed_bytes,
                    "media_type": media_type,
                    "page_number": None,  # DOCX doesn't have page numbers
                    "original_ext": ext,
                    "image_index": image_counter,  # Order of appearance
                    "total_images": None  # Will be set after all images extracted
                })
                image_counter += 1

            except Exception as e:
                print(f"Warning: Failed to extract image from DOCX: {str(e)}")
                continue

    # Set total_images count for all images
    for img in images:
        img['total_images'] = len(images)

    return images


def process_image_for_bedrock(image_bytes: bytes, image_ext: str) -> Tuple[bytes, str]:
    """
    Process and compress image to meet Bedrock requirements (max 3.75 MB).

    Args:
        image_bytes: Raw image bytes
        image_ext: Image extension (jpeg, png, gif, webp)

    Returns:
        Tuple of (processed image bytes, media type)
    """
    try:
        from PIL import Image
        import io
    except ImportError:
        print("Warning: Pillow not available for image processing. Using original image.")
        media_type = f"image/{image_ext if image_ext != 'jpg' else 'jpeg'}"
        return image_bytes, media_type

    MAX_SIZE_BYTES = 3_750_000  # 3.75 MB
    MAX_DIMENSION = 4096  # Max width/height for Bedrock

    # If image is already small enough, return as-is
    if len(image_bytes) <= MAX_SIZE_BYTES:
        media_type = f"image/{image_ext if image_ext != 'jpg' else 'jpeg'}"
        return image_bytes, media_type

    try:
        # Load image
        img = Image.open(io.BytesIO(image_bytes))

        # Convert RGBA to RGB if needed (for JPEG)
        if img.mode in ('RGBA', 'LA', 'P'):
            background = Image.new('RGB', img.size, (255, 255, 255))
            if img.mode == 'P':
                img = img.convert('RGBA')
            background.paste(img, mask=img.split()[-1] if img.mode in ('RGBA', 'LA') else None)
            img = background

        # Resize if too large
        if img.width > MAX_DIMENSION or img.height > MAX_DIMENSION:
            img.thumbnail((MAX_DIMENSION, MAX_DIMENSION), Image.Resampling.LANCZOS)

        # Compress image
        output = io.BytesIO()

        # Try JPEG compression first (smaller file size)
        quality = 85
        while quality > 20:
            output.seek(0)
            output.truncate()
            img.save(output, format='JPEG', quality=quality, optimize=True)

            if output.tell() <= MAX_SIZE_BYTES:
                break
            quality -= 10

        processed_bytes = output.getvalue()
        media_type = "image/jpeg"

        # If still too large, try more aggressive resize
        if len(processed_bytes) > MAX_SIZE_BYTES:
            scale_factor = 0.8
            while len(processed_bytes) > MAX_SIZE_BYTES and scale_factor > 0.3:
                new_width = int(img.width * scale_factor)
                new_height = int(img.height * scale_factor)
                resized_img = img.resize((new_width, new_height), Image.Resampling.LANCZOS)

                output.seek(0)
                output.truncate()
                resized_img.save(output, format='JPEG', quality=85, optimize=True)
                processed_bytes = output.getvalue()
                scale_factor -= 0.1

        return processed_bytes, media_type

    except Exception as e:
        print(f"Warning: Image processing failed: {str(e)}. Using original image.")
        media_type = f"image/{image_ext if image_ext != 'jpg' else 'jpeg'}"
        return image_bytes, media_type
